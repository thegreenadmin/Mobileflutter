#!/usr/bin/env python3
"""
Generate Microsoft Word document for Mobile App Build Pipeline
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_pipeline_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('Mobile App Build Pipeline', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Overview
    doc.add_heading('Overview', level=1)
    doc.add_paragraph(
        'This document illustrates the complete CI/CD pipeline flow from GitHub '
        'to both App Store (iOS) and Google Play Store (Android).'
    )
    
    # Pipeline Flow
    doc.add_heading('Pipeline Flow', level=1)
    doc.add_paragraph('The pipeline follows this sequence:')
    
    # Flow steps
    steps = [
        '1. Developer pushes code to GitHub or triggers manual workflow',
        '2. GitHub Actions triggers the build pipeline',
        '3. Platform selection: iOS, Android, or Both',
        '4. Parallel execution of iOS and Android build pipelines',
        '5. Code signing and environment setup',
        '6. Build based on environment (Beta or Release)',
        '7. Upload artifacts to respective platforms',
        '8. Auto-increment version numbers',
        '9. Commit and push version changes with [skip ci]',
        '10. Upload build artifacts to GitHub Actions',
        '11. Pipeline completion notification'
    ]
    
    for step in steps:
        doc.add_paragraph(step, style='List Number')
    
    # iOS Pipeline Section
    doc.add_heading('iOS Build Pipeline', level=1)
    
    doc.add_heading('iOS Build Process', level=2)
    ios_steps = [
        'Code Checkout: Fetch latest code from GitHub',
        'Environment Setup:',
        '  - Select Xcode',
        '  - Setup Ruby 3.2 + Fastlane',
        '  - Setup Flutter 3.41.6',
        '  - Install Flutter dependencies (flutter pub get)',
        '  - Install CocoaPods dependencies',
        'Code Signing:',
        '  - Use Fastlane Match for certificate management',
        '  - Sync signing certificates and provisioning profiles',
        '  - Required secrets: MATCH_PASSWORD, FASTLANE_PASSWORD, Apple API keys',
        'Build:',
        '  - Beta: Build and upload to TestFlight',
        '  - Release: Build for App Store distribution'
    ]
    
    for step in ios_steps:
        if step.startswith('  -'):
            doc.add_paragraph(step[3:], style='List Bullet 2')
        elif ':' in step and not step.startswith('  '):
            doc.add_paragraph(step.split(':')[0] + ':', style='List Bullet')
            doc.add_paragraph(step.split(':', 1)[1].strip())
        else:
            doc.add_paragraph(step, style='List Bullet')
    
    # Android Pipeline Section
    doc.add_heading('Android Build Pipeline', level=1)
    
    doc.add_heading('Android Build Process', level=2)
    android_steps = [
        'Code Checkout: Fetch latest code from GitHub',
        'Environment Setup:',
        '  - Setup Ruby 3.2 + Fastlane',
        '  - Setup Flutter 3.41.6',
        '  - Install Flutter dependencies',
        '  - Setup Java 17',
        'Code Signing:',
        '  - Decode base64-encoded keystore',
        '  - Required secrets: ANDROID_KEYSTORE_BASE64, ANDROID_KEYSTORE_PASSWORD, key alias',
        'Build:',
        '  - Beta: Build APK for testing',
        '  - Release: Build AAB for Play Store'
    ]
    
    for step in android_steps:
        if step.startswith('  -'):
            doc.add_paragraph(step[3:], style='List Bullet 2')
        elif ':' in step and not step.startswith('  '):
            doc.add_paragraph(step.split(':')[0] + ':', style='List Bullet')
            doc.add_paragraph(step.split(':', 1)[1].strip())
        else:
            doc.add_paragraph(step, style='List Bullet')
    
    # Step 2: Artifact Upload
    doc.add_heading('Step 2: Artifact Upload to Testing Platforms', level=1)
    
    doc.add_heading('iOS - TestFlight (Beta)', level=2)
    testflight_points = [
        'Automatically uploaded via Fastlane beta lane',
        'Available for internal and external testers',
        'Build artifacts stored for 30 days'
    ]
    for point in testflight_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_heading('Android - Google Play Store (Beta)', level=2)
    playstore_points = [
        'APK built for beta testing',
        'Uploaded via Fastlane',
        'Available through Play Store internal/alpha/beta tracks',
        'Build artifacts stored for 30 days'
    ]
    for point in playstore_points:
        doc.add_paragraph(point, style='List Bullet')
    
    # Step 3: Version Synchronization
    doc.add_heading('Step 3: Version Synchronization', level=1)
    
    doc.add_heading('iOS Version Management', level=2)
    ios_version_points = [
        'Build number auto-incremented in ios/Runner.xcodeproj/project.pbxproj',
        'Git commit with [skip ci] to prevent pipeline loops',
        'Ensures unique build for every commit'
    ]
    for point in ios_version_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_heading('Android Version Management', level=2)
    android_version_points = [
        'Version code auto-incremented in android/app/build.gradle',
        'Git commit with [skip ci] to prevent pipeline loops',
        'Ensures unique build for every commit'
    ]
    for point in android_version_points:
        doc.add_paragraph(point, style='List Bullet')
    
    # Triggers
    doc.add_heading('Triggers', level=1)
    
    doc.add_heading('Automatic Triggers', level=2)
    auto_triggers = [
        'Push to main branch',
        'Push to develop branch',
        'Push to devops/dev branch',
        'Changes in: lib/**, android/**, ios/**, pubspec.yaml'
    ]
    for trigger in auto_triggers:
        doc.add_paragraph(trigger, style='List Bullet')
    
    doc.add_heading('Manual Triggers', level=2)
    doc.add_paragraph('GitHub Actions workflow dispatch', style='List Bullet')
    doc.add_paragraph('Parameters:', style='List Bullet')
    doc.add_paragraph('Platform: ios, android, or both', style='List Bullet 2')
    doc.add_paragraph('Environment: beta or release', style='List Bullet 2')
    
    # Environments
    doc.add_heading('Environments', level=1)
    
    doc.add_heading('iOS Environments', level=2)
    doc.add_paragraph('APP_STORE_CONFIG: Used for both beta and release builds', style='List Bullet')
    doc.add_paragraph('Required variables: Apple Team ID, Fastlane user, Match Git URL, API keys', style='List Bullet')
    
    doc.add_heading('Android Environments', level=2)
    doc.add_paragraph('ANDROID_BETA_CONFIG: For beta/APK builds', style='List Bullet')
    doc.add_paragraph('ANDROID_RELEASE_CONFIG: For release/AAB builds', style='List Bullet')
    doc.add_paragraph('Required secrets: Keystore file, passwords, key alias', style='List Bullet')
    
    # Artifacts
    doc.add_heading('Artifacts', level=1)
    
    doc.add_heading('iOS Artifacts', level=2)
    doc.add_paragraph('.ipa file (30-day retention)', style='List Bullet')
    doc.add_paragraph('.dSYM.zip for crash reporting (30-day retention)', style='List Bullet')
    
    doc.add_heading('Android Artifacts', level=2)
    doc.add_paragraph('.aab file for Play Store (30-day retention)', style='List Bullet')
    doc.add_paragraph('.apk file for testing (30-day retention)', style='List Bullet')
    
    # Security
    doc.add_heading('Security', level=1)
    
    doc.add_heading('Secrets Management', level=2)
    security_points = [
        'All sensitive data stored in GitHub Secrets',
        'Certificates managed via Fastlane Match',
        'Keystore stored as base64-encoded secret',
        'API keys for App Store Connect and Google Play'
    ]
    for point in security_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_heading('Code Signing', level=2)
    signing_points = [
        'iOS: Match-based certificate management',
        'Android: Keystore-based signing',
        'Automatic sync on each build'
    ]
    for point in signing_points:
        doc.add_paragraph(point, style='List Bullet')
    
    # Component Interactions
    doc.add_heading('Component Interactions', level=1)
    
    doc.add_heading('iOS Interactions', level=2)
    ios_interactions = [
        'GitHub → macOS Runner: Job initiation and code checkout',
        'Fastlane ↔ Match: Certificate and provisioning profile retrieval',
        'Fastlane → App Store Connect: Authentication and API communication',
        'Fastlane → TestFlight: IPA upload for beta testing',
        'macOS Runner → Git: Build number increment commits',
        'macOS Runner → GitHub Actions: Artifact upload'
    ]
    for interaction in ios_interactions:
        doc.add_paragraph(interaction, style='List Number')
    
    doc.add_heading('Android Interactions', level=2)
    android_interactions = [
        'GitHub → Ubuntu Runner: Job initiation and code checkout',
        'Ubuntu Runner → Keystore: Local decoding and signing',
        'Fastlane → Google Play Console: APK/AAB upload via API',
        'Ubuntu Runner → Git: Version code increment commits',
        'Ubuntu Runner → GitHub Actions: Artifact upload'
    ]
    for interaction in android_interactions:
        doc.add_paragraph(interaction, style='List Number')
    
    doc.add_heading('Key Interaction Points', level=2)
    key_points = [
        'Secrets Flow: GitHub Secrets → Runner environment variables → Fastlane',
        'Code Signing: Match (iOS) / Keystore (Android) → Build process',
        'Version Control: Auto-increment → Git commit → GitHub push with [skip ci]',
        'Artifact Storage: Build outputs → GitHub Actions artifacts (30-day retention)',
        'Platform Communication: Fastlane acts as intermediary for Apple/Google APIs'
    ]
    for point in key_points:
        doc.add_paragraph(point, style='List Bullet')
    
    # Save document
    output_path = '/Users/yprattipatti/Mobileflutter/docs/Mobile_App_Build_Pipeline.docx'
    doc.save(output_path)
    print(f'Document created: {output_path}')

if __name__ == '__main__':
    try:
        create_pipeline_document()
    except ImportError:
        print('Error: python-docx library not found.')
        print('Install it with: pip install python-docx')
    except Exception as e:
        print(f'Error creating document: {e}')
